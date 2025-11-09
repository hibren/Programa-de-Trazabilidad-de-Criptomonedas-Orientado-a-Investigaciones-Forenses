import httpx
import base64
from typing import List
from app.database import reporte_collection
from app.models.reporte import Reporte
from datetime import datetime
import os
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from io import BytesIO
import pandas as pd

# 🔑 API Key de ChainAbuse
CHAINABUSE_API_KEY = "ca_QWQ2YkNGVmRBdDY2M3VQV3dLU2RSNFFHLlhCM2U5YjF3V1lrRi80K1hwYVNHbGc9PQ"

# 📁 Carpeta para guardar los reportes generados
GENERATED_PATH = "app/static/reportes"
os.makedirs(GENERATED_PATH, exist_ok=True)


# =====================================================
# 🔹 CONSULTAS DE REPORTES DESDE BD O API
# =====================================================

async def get_all_reportes() -> List[Reporte]:
    reportes_cursor = reporte_collection.find()
    reportes_list = await reportes_cursor.to_list(length=1000)
    return [Reporte(**reporte) for reporte in reportes_list]


async def fetch_reportes_by_address(address: str) -> List[Reporte]:
    if CHAINABUSE_API_KEY == "YOUR_CHAINABUSE_API_KEY":
        raise ValueError("La API key de ChainAbuse no ha sido configurada.")

    # 1️⃣ Buscar primero en la base de datos
    reportes_cursor = reporte_collection.find({"id_direccion": address})
    reportes_list = await reportes_cursor.to_list(length=1000)
    if reportes_list:
        print(f"✅ TRAIGO DESDE BD para la dirección {address}")
        return [Reporte(**reporte) for reporte in reportes_list]

    # 2️⃣ Si no existen, consultar la API externa
    print(f"🌐 TRAIGO DESDE API para la dirección {address}")
    auth_string = f"{CHAINABUSE_API_KEY}:"
    api_key_encoded = base64.b64encode(auth_string.encode()).decode()

    headers = {
        "accept": "application/json",
        "authorization": f"Basic {api_key_encoded}"
    }

    url = f"https://api.chainabuse.com/v0/reports?address={address}&includePrivate=false&page=1&perPage=50"

    async with httpx.AsyncClient() as client:
        try:
            response = await client.get(url, headers=headers)
            if response.status_code == 404:
                return []
            if response.status_code == 429:
                raise Exception("Has superado el límite de consultas a ChainAbuse. Intenta más tarde.")
            response.raise_for_status()
            data = response.json()
        except httpx.HTTPStatusError as e:
            raise Exception(f"API request failed with status {e.response.status_code}: {e.response.text}") from e
        except httpx.RequestError as e:
            raise Exception(f"API request failed: {e}") from e

    if "reports" not in data:
        return []

    # 3️⃣ Guardar en base lo que devuelve la API
    saved_reportes = []
    for report_data in data["reports"]:
        chainabuse_id = report_data["id"]

        created_at_str = report_data["createdAt"]
        if created_at_str.endswith('Z'):
            created_at_str = created_at_str[:-1] + '+00:00'

        try:
            created_at_dt = datetime.fromisoformat(created_at_str)
        except ValueError:
            created_at_dt = datetime.now()

        domains = [addr.get("domain") for addr in report_data.get("addresses", []) if addr.get("domain")]

        report_doc = {
            "chainabuse_id": chainabuse_id,
            "id_direccion": address,
            "scamCategory": report_data["scamCategory"],
            "createdAt": created_at_dt,
            "trusted": report_data.get("trusted", False),
            "domains": domains,
        }

        result = await reporte_collection.insert_one(report_doc)
        created_report = await reporte_collection.find_one({"_id": result.inserted_id})
        if created_report:
            saved_reportes.append(Reporte(**created_report))

    return saved_reportes


# 🔹 Nueva función: Obtener reportes sin consultar API
async def get_reportes_from_db_only(address: str) -> List[Reporte]:
    """
    Obtiene reportes SOLO de la base de datos, sin consultar la API.
    Útil para evitar rate limits.
    """
    reportes_cursor = reporte_collection.find({"id_direccion": address})
    reportes_list = await reportes_cursor.to_list(length=1000)
    return [Reporte(**reporte) for reporte in reportes_list]


# =====================================================
# 🔹 GENERACIÓN DE ARCHIVOS DE REPORTE
# =====================================================

def _crear_pdf(nombre_archivo: str, titulo: str, lineas: list[str]):
    path = os.path.join(GENERATED_PATH, nombre_archivo)
    c = canvas.Canvas(path, pagesize=A4)
    width, height = A4
    y = height - 80
    c.setFont("Helvetica-Bold", 16)
    c.drawString(60, y, titulo)
    c.setFont("Helvetica", 11)
    y -= 40
    for linea in lineas:
        c.drawString(60, y, str(linea))
        y -= 20
        if y < 50:
            c.showPage()
            c.setFont("Helvetica", 11)
            y = height - 80
    c.save()
    return path


# 🧩 Reporte de Riesgo por Dirección (PDF / Word / CSV)
async def generar_reporte_riesgo(address: str, formato: str = "PDF", force_api: bool = False) -> str:
    """
    Genera un reporte de riesgo para una dirección.
    
    Args:
        address: Dirección a analizar
        formato: PDF, WORD o CSV
        force_api: Si True, intenta consultar la API aunque falle. Si False, usa solo BD.
    """
    reportes = []
    api_error = None
    
    if force_api:
        # Modo: Intentar API (puede fallar con rate limit)
        try:
            print(f"🔄 Intentando obtener datos desde API...")
            reportes = await fetch_reportes_by_address(address)
        except Exception as e:
            api_error = str(e)
            print(f"⚠️ Error al consultar API: {api_error}")
            print(f"📂 Obteniendo datos solo de BD...")
            reportes = await get_reportes_from_db_only(address)
    else:
        # Modo: Solo BD (predeterminado, más seguro)
        print(f"📂 Obteniendo reportes solo de BD (sin consultar API)...")
        reportes = await get_reportes_from_db_only(address)
    
    # Construir contenido del reporte
    if not reportes:
        lineas = [
            f"Reporte de Riesgo para: {address}",
            "",
            "━" * 60,
            "RESULTADO DEL ANÁLISIS",
            "━" * 60,
            "",
            "✓ No se encontraron reportes de scam para esta dirección.",
            "",
            "Esto puede significar que:",
            "  • La dirección no ha sido reportada en ChainAbuse",
            "  • No hay datos disponibles en nuestra base de datos",
            "",
        ]
        if api_error and "límite" in api_error.lower():
            lineas.extend([
                "⚠ NOTA: No se pudo consultar la API de ChainAbuse debido a",
                "  límite de consultas. Se usaron solo datos en caché.",
                "",
            ])
    else:
        lineas = [
            f"Reporte de Riesgo para: {address}",
            "",
            "━" * 60,
            f"RESUMEN: {len(reportes)} reporte(s) encontrado(s)",
            "━" * 60,
            "",
        ]
        
        for idx, r in enumerate(reportes, 1):
            lineas.append(f"REPORTE #{idx}")
            lineas.append(f"  Categoría: {r.scamCategory}")
            lineas.append(f"  Fecha: {r.createdAt.strftime('%Y-%m-%d %H:%M')}")
            lineas.append(f"  Verificado: {'Sí' if r.trusted else 'No'}")
            if r.domains:
                lineas.append(f"  Dominios asociados: {', '.join(r.domains)}")
            lineas.append("")
        
        if api_error and "límite" in api_error.lower():
            lineas.extend([
                "━" * 60,
                "⚠ ADVERTENCIA",
                "━" * 60,
                "No se pudo verificar con la API de ChainAbuse debido a límite",
                "de consultas. Este reporte usa datos en caché de la BD.",
                "",
            ])
    
    # Agregar footer
    lineas.extend([
        "━" * 60,
        f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "Sistema de Trazabilidad Blockchain",
    ])

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    nombre_base = f"riesgo_{address[:12]}_{timestamp}"

    # === PDF ===
    if formato.upper() == "PDF":
        nombre = f"{nombre_base}.pdf"
        path = _crear_pdf(nombre, "Reporte de Riesgo por Dirección", lineas)

    # === WORD ===
    elif formato.upper() == "WORD":
        from docx import Document
        doc = Document()
        doc.add_heading("Reporte de Riesgo por Dirección", 0)
        for linea in lineas:
            doc.add_paragraph(linea)
        nombre = f"{nombre_base}.docx"
        path = os.path.join(GENERATED_PATH, nombre)
        doc.save(path)

    # === CSV ===
    elif formato.upper() == "CSV":
        data = []
        if reportes:
            for r in reportes:
                data.append({
                    "direccion": r.id_direccion,
                    "categoria": r.scamCategory,
                    "fecha": r.createdAt.strftime("%Y-%m-%d"),
                    "verificado": "Sí" if r.trusted else "No",
                    "dominios": ", ".join(r.domains or []),
                })
        else:
            data.append({
                "direccion": address,
                "categoria": "Sin reportes",
                "fecha": datetime.now().strftime("%Y-%m-%d"),
                "verificado": "N/A",
                "dominios": "N/A",
            })
        df = pd.DataFrame(data)
        nombre = f"{nombre_base}.csv"
        path = os.path.join(GENERATED_PATH, nombre)
        df.to_csv(path, index=False)

    else:
        raise ValueError(f"Formato '{formato}' no soportado. Usa PDF, WORD o CSV.")

    # 🧾 Guardar en la base de datos la referencia
    await reporte_collection.insert_one({
        "tipo": "riesgo",
        "id_direccion": address,
        "filename": os.path.basename(path),
        "path": path,
        "formato": formato.upper(),
        "createdAt": datetime.now(),
        "reportes_count": len(reportes),
        "api_error": api_error,
    })

    return path


# 📊 Reporte de Actividad por Período (CSV)
async def generar_reporte_actividad(fecha_inicio: str, fecha_fin: str) -> str:
    data = [
        {"fecha": fecha_inicio, "transacciones": 42, "volumen_btc": 0.87},
        {"fecha": fecha_fin, "transacciones": 37, "volumen_btc": 0.65},
    ]
    df = pd.DataFrame(data)
    nombre = f"actividad_{fecha_inicio}_{fecha_fin}.csv"
    path = os.path.join(GENERATED_PATH, nombre)
    df.to_csv(path, index=False)

    await reporte_collection.insert_one({
        "tipo": "actividad",
        "filename": nombre,
        "path": path,
        "createdAt": datetime.now()
    })
    return path


# 🕸️ Reporte de Clusters y Redes (PDF)
async def generar_reporte_clusters() -> str:
    lineas = [
        "Análisis de agrupaciones entre direcciones sospechosas.",
        "",
        "━" * 60,
        "CLUSTERS IDENTIFICADOS",
        "━" * 60,
        "",
        "• Cluster #1: 5 direcciones conectadas.",
        "• Cluster #2: 3 direcciones (2 reportadas en ChainAbuse).",
        "• Cluster #3: 8 direcciones vinculadas a dominio sin KYC.",
        "",
        "━" * 60,
        f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
    ]
    nombre = f"clusters_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    path = _crear_pdf(nombre, "Reporte de Clusters y Redes", lineas)

    await reporte_collection.insert_one({
        "tipo": "clusters",
        "filename": nombre,
        "path": path,
        "createdAt": datetime.now()
    })
    return path